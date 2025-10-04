from datetime import date, datetime, timezone
from calendar import monthrange
from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, flash, session, send_from_directory
import io
import zipfile
import base64
import os
import sqlite3
import secrets
import string
import requests
import tempfile
import subprocess
import shutil
from dotenv import load_dotenv
import re
from datetime import timedelta

try:
    from PIL import Image
except Exception:  # Pillow optional for image tools
    Image = None
try:
    import PyPDF2
except Exception:  # PyPDF2 optional for PDF tools
    PyPDF2 = None
try:
    import docx  # python-docx
except Exception:
    docx = None
try:
    import openpyxl
except Exception:
    openpyxl = None
try:
    import qrcode
except Exception:  # qrcode optional for QR tool
    qrcode = None
try:
    from spellchecker import SpellChecker
except Exception:
    SpellChecker = None
try:
    import speedtest  # speedtest-cli
except Exception:
    speedtest = None

# Load environment variables
load_dotenv()

app = Flask(__name__, static_folder="static", template_folder="templates", static_url_path='/static')
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'dev-secret-key-change-in-production')

DB_PATH = os.path.join(os.path.dirname(__file__), "data.db")

def init_db() -> None:
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True) if os.path.dirname(DB_PATH) else None
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS short_urls (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                code TEXT UNIQUE NOT NULL,
                url TEXT NOT NULL,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
            """
        )

def generate_code(length: int = 7) -> str:
    alphabet = string.ascii_letters + string.digits
    return "".join(secrets.choice(alphabet) for _ in range(length))

def difference_ymd(start_date: date, end_date: date) -> tuple[int, int, int]:
    years = end_date.year - start_date.year
    months = end_date.month - start_date.month
    days = end_date.day - start_date.day

    if days < 0:
        borrow_month = end_date.month - 1 or 12
        borrow_year = end_date.year if end_date.month != 1 else end_date.year - 1
        days_in_borrow_month = monthrange(borrow_year, borrow_month)[1]
        days += days_in_borrow_month
        months -= 1

    if months < 0:
        months += 12
        years -= 1

    return years, months, days

def add_months(base_date: date, months_to_add: int) -> date:
    total_months = base_date.year * 12 + (base_date.month - 1) + months_to_add
    new_year = total_months // 12
    new_month = total_months % 12 + 1
    last_day = monthrange(new_year, new_month)[1]
    new_day = min(base_date.day, last_day)
    return date(new_year, new_month, new_day)

def months_and_extra_days(start_date: date, end_date: date) -> tuple[int, int]:
    months_total = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
    if end_date.day < start_date.day:
        months_total -= 1
    anchor = add_months(start_date, months_total)
    extra_days = (end_date - anchor).days
    return months_total, extra_days

# Main Routes
@app.route("/")
def home():
    return render_template("home.html", current_user=None)

@app.get("/all-tools")
def all_tools():
    return render_template("all_tools.html", current_user=None)

@app.get("/age")
def age_calculator():
    return render_template("age.html", current_user=None)

@app.post("/api/diff")
def api_diff():
    try:
        data = request.get_json(force=True)
        start_input = data.get("startDate")
        end_input = data.get("endDate")
        if not start_input or not end_input:
            return jsonify({"error": "startDate and endDate are required (YYYY-MM-DD)"}), 400
        start_parts = list(map(int, start_input.split("-")))
        end_parts = list(map(int, end_input.split("-")))
        start_date = date(*start_parts)
        end_date = date(*end_parts)
    except Exception as ex:
        return jsonify({"error": f"Invalid input: {ex}"}), 400

    if end_date < start_date:
        start_date, end_date = end_date, start_date

    years, months, days = difference_ymd(start_date, end_date)
    total_months, extra_days = months_and_extra_days(start_date, end_date)
    total_days = (end_date - start_date).days

    return jsonify({
        "years": years,
        "months": months,
        "days": days,
        "totalMonths": total_months,
        "extraDays": extra_days,
        "totalDays": total_days
    })

# ---------------------------- PDF Tools ----------------------------

@app.get("/pdf")
def pdf_tools_page():
    return render_template("pdf_tools.html", current_user=None)

@app.post("/api/pdf/merge")
def api_pdf_merge():
    if PyPDF2 is None:
        return jsonify({"error": "PyPDF2 not installed"}), 500
    files = request.files.getlist("files")
    if not files:
        return jsonify({"error": "No PDF files uploaded"}), 400
    merger = PyPDF2.PdfMerger()
    for f in files:
        merger.append(io.BytesIO(f.read()))
    output = io.BytesIO()
    merger.write(output)
    merger.close()
    output.seek(0)
    return send_file(output, mimetype="application/pdf", as_attachment=True, download_name="merged.pdf")

@app.post("/api/pdf/split")
def api_pdf_split():
    if PyPDF2 is None:
        return jsonify({"error": "PyPDF2 not installed"}), 500
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No PDF uploaded"}), 400
    reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for i in range(len(reader.pages)):
            writer = PyPDF2.PdfWriter()
            writer.add_page(reader.pages[i])
            single_buffer = io.BytesIO()
            writer.write(single_buffer)
            zf.writestr(f"page_{i+1}.pdf", single_buffer.getvalue())
    zip_buffer.seek(0)
    return send_file(zip_buffer, mimetype="application/zip", as_attachment=True, download_name="split_pages.zip")

@app.post("/api/pdf/compress")
def api_pdf_compress():
    if PyPDF2 is None:
        return jsonify({"error": "PyPDF2 not installed"}), 500
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No PDF uploaded"}), 400
    reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
    writer = PyPDF2.PdfWriter()
    for page in reader.pages:
        try:
            page.compress_content_streams()
        except Exception:
            pass
        writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    out.seek(0)
    return send_file(out, mimetype="application/pdf", as_attachment=True, download_name="compressed.pdf")

@app.post("/api/pdf/to-word")
def api_pdf_to_word():
    if PyPDF2 is None or docx is None:
        return jsonify({"error": "PyPDF2 and python-docx required"}), 500
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No PDF uploaded"}), 400
    reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
    document = docx.Document()
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        document.add_heading(f"Page {i}", level=2)
        for line in text.splitlines():
            document.add_paragraph(line)
    out = io.BytesIO()
    document.save(out)
    out.seek(0)
    return send_file(out, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", as_attachment=True, download_name="converted.docx")

@app.post("/api/pdf/to-excel")
def api_pdf_to_excel():
    if PyPDF2 is None or openpyxl is None:
        return jsonify({"error": "PyPDF2 and openpyxl required"}), 500
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No PDF uploaded"}), 400
    reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "PDF Text"
    row_idx = 1
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").splitlines()
        ws.cell(row=row_idx, column=1, value=f"Page {i}")
        row_idx += 1
        for line in text:
            ws.cell(row=row_idx, column=1, value=line)
            row_idx += 1
        row_idx += 1
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return send_file(out, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", as_attachment=True, download_name="converted.xlsx")

# ------------------------- File Converter -------------------------

@app.get("/file-converter")
def file_converter_page():
    return render_template("file_converter.html", current_user=None)

@app.post("/api/convert/image")
def api_convert_image():
    if Image is None:
        return jsonify({"error": "Pillow not installed"}), 500
    file = request.files.get("file")
    target = request.form.get("target", "png").lower()
    if not file:
        return jsonify({"error": "No image uploaded"}), 400
    image = Image.open(file.stream).convert("RGB")
    output = io.BytesIO()
    image.save(output, format=target.upper())
    output.seek(0)
    return send_file(output, mimetype=f"image/{target}", as_attachment=True, download_name=f"converted.{target}")

@app.post("/api/convert/video")
def api_convert_video():
    # Requires ffmpeg installed on OS and present in PATH
    file = request.files.get("file")
    target = request.form.get("target", "mp4").lower()
    if not file:
        return jsonify({"error": "No video uploaded"}), 400
    if shutil.which("ffmpeg") is None:
        return jsonify({"error": "ffmpeg not found. Please install ffmpeg and add to PATH"}), 500
    with tempfile.TemporaryDirectory() as td:
        in_path = os.path.join(td, "input")
        out_path = os.path.join(td, f"output.{target}")
        file.save(in_path)
        cmd = ["ffmpeg", "-y", "-i", in_path, out_path]
        try:
            subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        except subprocess.CalledProcessError:
            return jsonify({"error": "Conversion failed"}), 500
        return send_file(out_path, as_attachment=True, download_name=f"converted.{target}")

# ------------------------- URL Shortener -------------------------

@app.get("/shortener")
def url_shortener_page():
    return render_template("url_shortener.html", current_user=None)

@app.post("/api/shorten")
def api_shorten():
    data = request.get_json(force=True)
    url = data.get("url")
    if not url:
        return jsonify({"error": "url is required"}), 400
    code = generate_code()
    with sqlite3.connect(DB_PATH) as conn:
        try:
            conn.execute("INSERT INTO short_urls(code, url) VALUES(?, ?)", (code, url))
        except sqlite3.IntegrityError:
            code = generate_code()
            conn.execute("INSERT INTO short_urls(code, url) VALUES(?, ?)", (code, url))
    short = request.host_url.rstrip("/") + url_for("redirect_short", code=code)
    return jsonify({"code": code, "shortUrl": short})

@app.get("/u/<code>")
def redirect_short(code: str):
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute("SELECT url FROM short_urls WHERE code = ?", (code,))
        row = cur.fetchone()
    if not row:
        return render_template("url_shortener.html", error="Invalid code"), 404
    return redirect(row[0])

# ----------------------- QR Code Tools ---------------------------

@app.get("/qr")
def qr_tools_page():
    return render_template("qr_tools.html", current_user=None)

@app.post("/api/qr/generate")
def api_qr_generate():
    if qrcode is None:
        return jsonify({"error": "qrcode not installed"}), 500
    data = request.get_json(force=True)
    text = data.get("text", "")
    size = int(data.get("size", 256))
    fill = data.get("fill", "#000000")
    back = data.get("back", "#FFFFFF")
    if not text:
        return jsonify({"error": "text is required"}), 400
    qr = qrcode.QRCode(version=None, box_size=10, border=2)
    qr.add_data(text)
    qr.make(fit=True)
    img = qr.make_image(fill_color=fill, back_color=back).resize((size, size))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    b64 = base64.b64encode(buf.getvalue()).decode("ascii")
    return jsonify({"pngBase64": f"data:image/png;base64,{b64}"})

# ---------------- Image Resizer & Compressor (Bulk) --------------

@app.get("/image-tools")
def image_tools_page():
    return render_template("image_tools.html", current_user=None)

@app.post("/api/image/bulk")
def api_image_bulk():
    if Image is None:
        return jsonify({"error": "Pillow not installed"}), 500
    files = request.files.getlist("files")
    width = request.form.get("width")
    height = request.form.get("height")
    quality = int(request.form.get("quality", 80))
    fmt = request.form.get("format", "jpeg").lower()
    if not files:
        return jsonify({"error": "No images uploaded"}), 400
    width_i = int(width) if width else None
    height_i = int(height) if height else None
    mem_zip = io.BytesIO()
    with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in files:
            img = Image.open(f.stream)
            img = img.convert("RGB")
            if width_i or height_i:
                img = img.resize((width_i or img.width, height_i or img.height))
            out = io.BytesIO()
            save_kwargs = {"quality": quality}
            if fmt == "png":
                save_kwargs.pop("quality", None)
            img.save(out, format=fmt.upper(), **save_kwargs)
            zf.writestr(os.path.splitext(f.filename)[0] + f"_processed.{fmt}", out.getvalue())
    mem_zip.seek(0)
    return send_file(mem_zip, mimetype="application/zip", as_attachment=True, download_name="images_processed.zip")

# ------------------------ Unit Converter -------------------------

@app.get("/unit-converter")
def unit_converter_page():
    return render_template("unit_converter.html", current_user=None)

# ------------------------ Internet Speed Test --------------------

@app.get("/speed-test")
def speed_test_page():
    return render_template("speed_test.html", current_user=None)

@app.post("/api/speedtest")
def api_speedtest():
    if speedtest is None:
        return jsonify({"error": "speedtest-cli not installed"}), 500
    s = speedtest.Speedtest()
    s.get_best_server()
    down = s.download() / 1_000_000
    up = s.upload() / 1_000_000
    ping = s.results.ping
    return jsonify({"downloadMbps": round(down, 2), "uploadMbps": round(up, 2), "pingMs": round(ping, 2)})

# ------------------------ Currency Converter ---------------------

@app.post("/api/currency")
def api_currency():
    try:
        data = request.get_json(force=True)
        amount = float(data.get("amount"))
        from_cur = data.get("from").upper().strip()
        to_cur = data.get("to").upper().strip()
        if not from_cur or not to_cur:
            return jsonify({"error": "from and to currencies required"}), 400
    except Exception:
        return jsonify({"error": "Invalid input"}), 400
    try:
        r = requests.get("https://api.exchangerate.host/convert", params={"from": from_cur, "to": to_cur, "amount": amount}, timeout=15)
        j = r.json()
        if not r.ok or "result" not in j:
            return jsonify({"error": "Rate lookup failed"}), 502
        return jsonify({"amount": amount, "from": from_cur, "to": to_cur, "converted": j["result"]})
    except Exception:
        return jsonify({"error": "Conversion service unreachable"}), 502

# ------------------------ Screen Recorder (Web) ------------------

@app.get("/screen-recorder")
def screen_recorder_page():
    return render_template("screen_recorder.html", current_user=None)

# ---------------------- Grammar / Spell Checker ------------------

@app.get("/grammar")
def grammar_page():
    return render_template("grammar.html", current_user=None)

@app.post("/api/grammar")
def api_grammar():
    text = request.get_json(force=True).get("text", "")
    if not text:
        return jsonify({"errors": []})
    results = []
    if SpellChecker is not None:
        sp = SpellChecker()
        words = [w.strip(".,!?;:\"'()[]{}") for w in text.split()]
        misspelled = sp.unknown([w for w in words if w])
        for w in misspelled:
            suggestions = list(sp.candidates(w))[:5]
            results.append({"word": w, "suggestions": suggestions})
    return jsonify({"errors": results})

# ------------------------ Footer Links -------------------------

@app.get("/about")
def about_page():
    return render_template("about.html", current_user=None)

@app.get("/contact")
def contact_page():
    return render_template("contact.html", current_user=None)

@app.get("/privacy")
def privacy_page():
    return render_template("privacy.html", current_user=None)

@app.get("/terms")
def terms_page():
    return render_template("terms.html", current_user=None)

# ------------------------ SEO Routes -------------------------

@app.route('/sitemap.xml')
def sitemap():
    return send_from_directory('.', 'sitemap.xml', mimetype='application/xml')

@app.route('/robots.txt')
def robots():
    return send_from_directory('.', 'robots.txt', mimetype='text/plain')

if __name__ == "__main__":
    init_db()
    print("🚀 Starting Toolflock - All tools should be working!")
    print("📱 Access the website at: http://127.0.0.1:5000")
    app.run(debug=True, host='127.0.0.1', port=5000)